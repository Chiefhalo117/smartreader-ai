from flask import Flask, render_template, request, jsonify
import openai
from openai import OpenAI
import os
from datetime import datetime
import logging
from dotenv import load_dotenv
import re
from collections import defaultdict
import time

# 在现有导入之后添加

# 全局使用计数器
daily_usage = defaultdict(lambda: {'count': 0, 'date': time.strftime('%Y-%m-%d')})
global_api_calls = {'count': 0, 'date': time.strftime('%Y-%m-%d')}

load_dotenv()

# 配置日志
logging.basicConfig(level=logging.INFO)

app = Flask(__name__)

# OpenAI配置
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

def check_usage_limit(request):
    """检查用户每日使用限制和全局API限制"""
    user_ip = request.remote_addr
    today = time.strftime('%Y-%m-%d')
    
    # 检查全局API调用限制
    if global_api_calls['date'] != today:
        global_api_calls['count'] = 0
        global_api_calls['date'] = today
    
    # 全局限制：每天1000次API调用
    if global_api_calls['count'] >= 1000:
        return False, 0, "服务今日使用量已达上限，请明天再试。感谢您的理解！🙏"
    
    # 用户个人限制
    if daily_usage[user_ip]['date'] != today:
        daily_usage[user_ip] = {'count': 0, 'date': today}
    
    if daily_usage[user_ip]['count'] >= 10:
        return False, daily_usage[user_ip]['count'], f'您今天已使用{daily_usage[user_ip]["count"]}次，达到每日10次限制。明天再来试试吧！💪'
    
    # 增加计数
    daily_usage[user_ip]['count'] += 1
    global_api_calls['count'] += 3  # 每次分析调用3次API
    
    return True, daily_usage[user_ip]['count'], None

@app.route('/')
def index():
    """主页面路由"""
    return render_template('index.html')

@app.route('/api/analyze', methods=['POST'])
def analyze_text():
    """文本分析API - 分别调用AI三次"""
    try:
        # 检查使用限制
        can_use, usage_count, error_msg = check_usage_limit(request)
        if not can_use:
            return jsonify({
                'error': error_msg,
                'usage_limit_reached': True
            }), 429
        
        # 继续现有逻辑...
        data = request.get_json()
        
        role = data.get('role', 'reader')
        goal = data.get('goal', 'understand')
        language = data.get('language', 'English')
        text = data.get('text', '')

        if not text.strip():
            return jsonify({'error': 'Text is required'}), 400
        
        # 检查是否是URL
        url_pattern = r'^https?:\/\/.+'
        if re.match(url_pattern, text.strip()):
            # 如果是URL，先获取内容
            try:
                import requests
                from bs4 import BeautifulSoup
                
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                
                response = requests.get(text.strip(), headers=headers, timeout=10)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.content, 'html.parser')
                for script in soup(["script", "style"]):
                    script.decompose()
                
                page_text = soup.get_text()
                lines = (line.strip() for line in page_text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                if len(text) > 10000:
                    text = text[:10000] + "..."
                    
            except Exception as e:
                logging.error(f"Error fetching URL: {str(e)}")
                
                # 友好的错误信息
                error_msg = str(e).lower()
                
                if 'timeout' in error_msg:
                    user_msg = "网页加载超时，请检查URL或稍后重试。"
                elif 'connection' in error_msg:
                    user_msg = "无法连接到该网站，请检查URL是否正确。"
                elif '404' in error_msg:
                    user_msg = "网页不存在(404错误)，请检查URL是否正确。"
                elif '403' in error_msg or 'forbidden' in error_msg:
                    user_msg = "网站拒绝访问，可能需要登录。请复制内容到文本框分析。"
                elif 'ssl' in error_msg or 'certificate' in error_msg:
                    user_msg = "网站安全证书问题，请尝试复制内容到文本框分析。"
                else:
                    user_msg = "无法获取网页内容，请检查URL或直接复制内容到文本框。"
                
                return jsonify({'error': user_msg}), 500
        
        # 分别生成三个提示词
        summary_prompt = create_summary_prompt(role, goal, language, text)
        takeaway_prompt = create_takeaway_prompt(role, goal, language, text)
        question_prompt = create_question_prompt(role, goal, language, text)
        
        # 调用AI三次
        summary_response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": summary_prompt}],
            max_tokens=800,
            temperature=0.7,
        )
        
        takeaway_response = client.chat.completions.create(
            model="gpt-3.5-turbo", 
            messages=[{"role": "user", "content": takeaway_prompt}],
            max_tokens=800,
            temperature=0.7,
        )
        
        question_response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": question_prompt}],
            max_tokens=800,
            temperature=0.7,
        )
        
        # 在返回结果中添加使用统计
        result = {
            'summary': summary_response.choices[0].message.content.strip(),
            'takeaways': takeaway_response.choices[0].message.content.strip(),
            'questions': question_response.choices[0].message.content.strip(),
            'usage_info': {
                'used_today': usage_count,
                'user_limit': 10,
                'user_remaining': 10 - usage_count,
                'global_used': global_api_calls['count'],
                'global_limit': 1000,
                'global_remaining': 1000 - global_api_calls['count']
            }
        }

        logging.info(f"Analysis completed. User usage: {usage_count}/10, Global API calls: {global_api_calls['count']}/1000")
        return jsonify(result)
        
        logging.info(f"Analysis completed for role: {role}, goal: {goal}")
        # 在返回结果中添加使用统计
        result = {
            'summary': summary_response.choices[0].message.content.strip(),
            'takeaways': takeaway_response.choices[0].message.content.strip(),
            'questions': question_response.choices[0].message.content.strip(),
            'usage_info': {
                'used_today': usage_count,
                'limit': 10,
                'remaining': 10 - usage_count
            }
        }

        logging.info(f"Analysis completed for role: {role}, goal: {goal}, usage: {usage_count}/10")
        return jsonify(result)
        
    except Exception as e:
        logging.error(f"Error in analyze_text: {str(e)}")
        
        # 友好的错误信息
        error_msg = str(e).lower()
        
        if 'rate limit' in error_msg:
            user_msg = "AI服务器忙碌中，请等待30秒后重试。"
        elif 'timeout' in error_msg:
            user_msg = "网络连接超时，请检查网络后重试。"
        elif 'context length' in error_msg or 'too long' in error_msg:
            user_msg = "内容太长了！请缩短到5000字以内，或分段分析。"
        elif 'invalid api key' in error_msg:
            user_msg = "服务暂时不可用，请稍后重试。"
        else:
            user_msg = "分析失败，请尝试刷新页面重试，或联系客服。"
        
        return jsonify({'error': user_msg}), 500

@app.route('/api/analyze-file', methods=['POST'])
def analyze_file():
    """直接分析文件API"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        role = request.form.get('role', 'reader')
        goal = request.form.get('goal', 'understand')
        language = request.form.get('language', 'English')
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        filename = file.filename.lower()
        
        try:
            if filename.endswith('.pdf'):
                import PyPDF2
                import io
                
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                    
            elif filename.endswith('.txt'):
                text = file.read().decode('utf-8')
                
            elif filename.endswith(('.doc', '.docx')):
                from docx import Document
                import io
                
                doc = Document(io.BytesIO(file.read()))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                    
            else:
                return jsonify({'error': 'Unsupported file type'}), 400
            
            if len(text.strip()) < 50:
                return jsonify({'error': 'File content too short to analyze'}), 400
            
            # 限制文本长度
            if len(text) > 15000:
                text = text[:15000] + "..."
            
            # 分别生成三个提示词
            summary_prompt = create_summary_prompt(role, goal, language, text)
            takeaway_prompt = create_takeaway_prompt(role, goal, language, text)
            question_prompt = create_question_prompt(role, goal, language, text)

            # 打印提示词检查
            print("=== TAKEAWAY PROMPT ===")
            print(takeaway_prompt)
            print("=== QUESTION PROMPT ===")
            print(question_prompt)
            
            # 调用AI三次
            summary_response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": summary_prompt}],
                max_tokens=800,
                temperature=0.7,
            )
            
            takeaway_response = client.chat.completions.create(
                model="gpt-3.5-turbo", 
                messages=[{"role": "user", "content": takeaway_prompt}],
                max_tokens=800,
                temperature=0.7,
            )
            
            question_response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": question_prompt}],
                max_tokens=800,
                temperature=0.7,
            )
            
            result = {
                'summary': summary_response.choices[0].message.content.strip(),
                'takeaways': takeaway_response.choices[0].message.content.strip(),
                'questions': question_response.choices[0].message.content.strip()
            }
            
            logging.info(f"File analysis completed for: {file.filename}")
            return jsonify(result)
            
        except Exception as e:
            logging.error(f"Error processing file: {str(e)}")
            
            # 友好的错误信息
            error_msg = str(e).lower()
            
            if 'pdf' in error_msg and ('decrypt' in error_msg or 'password' in error_msg):
                user_msg = "PDF文件有密码保护，请提供无密码的文件。"
            elif 'memory' in error_msg or 'too large' in error_msg:
                user_msg = "文件太大！请上传小于5MB的文件，或复制内容到文本框。"
            elif 'corrupt' in error_msg or 'invalid' in error_msg:
                user_msg = "文件可能已损坏，请重新保存文件或使用其他格式。"
            elif 'docx' in error_msg or 'doc' in error_msg:
                user_msg = "Word文件处理失败，请尝试另存为PDF格式或复制内容到文本框。"
            else:
                user_msg = "文件处理失败，请尝试：1) 转换为PDF或TXT格式 2) 复制内容到文本框"
            
            return jsonify({'error': user_msg}), 500
            
    except Exception as e:
        logging.error(f"Error in analyze_file: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    """用户反馈API - 发送到Google Form"""
    try:
        data = request.get_json()
        feedback = data.get('feedback', '')
        payment = data.get('payment', '')
        role = data.get('role', '')
        goal = data.get('goal', '')
        language = data.get('language', '')
        
        if not feedback.strip():
            return jsonify({'error': 'Feedback is required'}), 400
        
        # Google Form 提交函数
        success = submit_to_google_form(feedback, payment, role, goal, language)
        
        if success:
            # 记录到本地日志作为备份
            feedback_data = {
                'feedback': feedback,
                'payment': payment,
                'role': role,
                'goal': goal,
                'language': language,
                'timestamp': datetime.now().isoformat()
            }
            
            logging.info(f"Feedback submitted to Google Form: {feedback_data}")
            
            return jsonify({
                'success': True, 
                'message': 'Thank you for your feedback!'
            })
        else:
            # 如果 Google Form 提交失败，至少保存到本地日志
            logging.error(f"Failed to submit to Google Form, saving locally: {feedback}")
            return jsonify({
                'success': True,  # 对用户来说仍然成功
                'message': 'Feedback received! Thank you!'
            })
        
    except Exception as e:
        logging.error(f"Error in submit_feedback: {str(e)}")
        return jsonify({'error': 'Failed to submit feedback'}), 500

def submit_to_google_form(feedback, payment, role, goal, language):
    """提交反馈到 Google Form"""
    try:
        import requests
        import urllib.parse
        
        # 构建GET请求的URL
        base_url = "https://docs.google.com/forms/d/e/1FAIpQLSd47a59Ym5cl9U68DxIlJ3HzZmwmAMQKncUcvWtbbvsfvMnXQ/formResponse"
        
        # 构建查询参数
        # 映射前端值到表单期望的值
        role_mapping = {
            'student': 'Student',
            'teacher': 'Teacher', 
            'sales': 'Sales',
            'customer': 'Customer',
            'product manager': 'Product Manager',
            'analyst': 'Analyst',
            'supervisor': 'Supervisor',
            'reader': 'Reader',
            'consultant': 'Consultant',
            'developer': 'Developer',
            'marketer': 'Marketer'
        }

        goal_mapping = {
            'understand': 'Understand the key ideas',
            'explain': 'Be able to explain it to others',
            'apply': 'Apply it to my work',
            'argue': 'Critically challenge or argue with it'
        }

        language_mapping = {
            'English': 'English',
            'Chinese': 'Chinese',
            'Spanish': 'Spanish'
        }

        # 转换值
        mapped_role = role_mapping.get(role, role)
        mapped_goal = goal_mapping.get(goal, goal)
        mapped_language = language_mapping.get(language, language)

        params = {
            'entry.93085': feedback,
            'entry.1377510532': payment,
            'entry.143639852': mapped_role,
            'entry.1422681914': mapped_goal,
            'entry.1210374444': mapped_language,
            'entry.1729421386': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'submit': 'Submit'
        }
                
        # 使用GET请求
        response = requests.get(
            base_url,
            params=params,
            timeout=10,
            headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            },
            allow_redirects=True
        )
        
        # 添加调试信息
        logging.info(f"GET Request URL: {response.url}")
        logging.info(f"Response status code: {response.status_code}")
        logging.info(f"Response text preview: {response.text[:300]}")
        
        # 检查响应中是否包含成功提示
        success_indicators = [
            'submitted', 'thank you', 'received', '已提交', '感谢', 
            'your response has been recorded', 'response recorded',
            'smartreader ai 用户反馈',  # 你的表单标题
            'form submitted', 'submission recorded'
        ]
        response_text_lower = response.text.lower()

        # 如果状态码是200且响应包含表单标题，通常表示成功
        is_success = (
            response.status_code == 200 and 
            (any(indicator in response_text_lower for indicator in success_indicators) or
            'smartreader ai' in response_text_lower)  # 表单标题出现通常表示成功
        )

        logging.info(f"Response contains form title: {'smartreader ai' in response_text_lower}")
        
        logging.info(f"Form submission success: {is_success}")
        return is_success
        
    except Exception as e:
        logging.error(f"Error submitting to Google Form: {str(e)}")
        return False


@app.route('/api/fetch-url', methods=['POST'])
def fetch_url():
    """获取URL内容API"""
    try:
        import requests
        from bs4 import BeautifulSoup
        
        data = request.get_json()
        url = data.get('url', '')
        
        if not url:
            return jsonify({'error': 'URL is required'}), 400
        
        # 发送HTTP请求获取网页内容
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # 解析HTML并提取文本
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # 移除脚本和样式元素
        for script in soup(["script", "style"]):
            script.decompose()
        
        # 获取文本内容
        text = soup.get_text()
        
        # 清理文本
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        # 限制文本长度
        if len(text) > 10000:
            text = text[:10000] + "..."
        
        logging.info(f"Successfully fetched content from URL: {url}")
        return jsonify({'content': text})
        
    except Exception as e:
        logging.error(f"Error fetching URL: {str(e)}")
        
        # 友好的错误信息
        error_msg = str(e).lower()
        
        if 'timeout' in error_msg:
            user_msg = "网页加载超时，请检查URL或稍后重试。"
        elif 'connection' in error_msg:
            user_msg = "无法连接到该网站，请检查URL是否正确。"
        elif '404' in error_msg:
            user_msg = "网页不存在(404错误)，请检查URL是否正确。"
        elif '403' in error_msg or 'forbidden' in error_msg:
            user_msg = "网站拒绝访问，可能需要登录。请复制内容到文本框分析。"
        elif 'ssl' in error_msg or 'certificate' in error_msg:
            user_msg = "网站安全证书问题，请尝试复制内容到文本框分析。"
        else:
            user_msg = "无法获取网页内容，请检查URL或直接复制内容到文本框。"
        
        return jsonify({'error': user_msg}), 500

@app.route('/api/process-file', methods=['POST'])
def process_file():
    """处理文件上传API"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        filename = file.filename.lower()
        
        try:
            if filename.endswith('.pdf'):
                import PyPDF2
                import io
                
                # 读取PDF内容
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                    
            elif filename.endswith('.txt'):
                # 读取文本文件
                text = file.read().decode('utf-8')
                
            elif filename.endswith(('.doc', '.docx')):
                from docx import Document
                import io
                
                doc = Document(io.BytesIO(file.read()))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                    
            else:
                return jsonify({'error': 'Unsupported file type'}), 400
            
            # 限制文本长度
            if len(text) > 15000:
                text = text[:15000] + "..."
                
            return jsonify({'content': text.strip()})
            
        except Exception as e:
            logging.error(f"Error processing file: {str(e)}")
            
            # 友好的错误信息
            error_msg = str(e).lower()
            
            if 'pdf' in error_msg and ('decrypt' in error_msg or 'password' in error_msg):
                user_msg = "PDF文件有密码保护，请提供无密码的文件。"
            elif 'memory' in error_msg or 'too large' in error_msg:
                user_msg = "文件太大！请上传小于5MB的文件，或复制内容到文本框。"
            elif 'corrupt' in error_msg or 'invalid' in error_msg:
                user_msg = "文件可能已损坏，请重新保存文件或使用其他格式。"
            elif 'docx' in error_msg or 'doc' in error_msg:
                user_msg = "Word文件处理失败，请尝试另存为PDF格式或复制内容到文本框。"
            else:
                user_msg = "文件处理失败，请尝试：1) 转换为PDF或TXT格式 2) 复制内容到文本框"
            
            return jsonify({'error': user_msg}), 500
            
    except Exception as e:
        logging.error(f"Error in process_file: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500

def create_summary_prompt(role, goal, language, text):
    """创建摘要提示词"""
    
    role_contexts = {
        'student': 'As a student',
        'teacher': 'As a teacher',
        'sales': 'As a sales professional',
        'customer': 'As a customer',
        'product manager': 'As a product manager',
        'analyst': 'As an analyst',
        'supervisor': 'As a supervisor',
        'reader': 'As a reader',
        'consultant': 'As a consultant',
        'developer': 'As a developer',
        'marketer': 'As a marketer'
    }
    
    goal_contexts = {
        'understand': 'to understand the key ideas',
        'explain': 'to explain it to others',
        'apply': 'to apply it to work',
        'argue': 'to critically analyze it'
    }
    
    role_context = role_contexts.get(role, 'As a reader')
    goal_context = goal_contexts.get(goal, 'to understand the key ideas')
    
    word_count = len(text.split())
    if word_count < 300:
        point_count = "2-3"
    elif word_count < 1000:
        point_count = "3-4"
    else:
        point_count = "4-6"
    
    # 添加强制语言指令
    if language.lower() == 'chinese':
        language_instruction = "You must respond completely in Chinese (中文). Do not use any English words."
    elif language.lower() == 'spanish':
        language_instruction = "You must respond completely in Spanish (Español). Do not use any English words."
    else:
        language_instruction = "You must respond completely in English."
    
    return f"""
{role_context} who wants {goal_context}, provide {point_count} key summary points of the following content.

{language_instruction}

Content:
{text}

Requirements:
- Write entirely in {language}
- Must use numbered format (1. 2. 3.)
- Each point should start with a bold theme: **Theme**: explanation
- Focus on the main facts and key information
- Be specific and substantive
- No introductory text, just the numbered points
- CRITICAL: Your entire response must be in {language} language only
"""

def create_takeaway_prompt(role, goal, language, text):
    """创建要点提示词"""
    
    role_contexts = {
        'student': 'As a student',
        'teacher': 'As a teacher', 
        'sales': 'As a sales professional',
        'customer': 'As a customer',
        'product manager': 'As a product manager',
        'analyst': 'As an analyst',
        'supervisor': 'As a supervisor',
        'reader': 'As a reader',
        'consultant': 'As a consultant',
        'developer': 'As a developer',
        'marketer': 'As a marketer'
    }
    
    role_context = role_contexts.get(role, 'As a reader')
    
    word_count = len(text.split())
    if word_count < 300:
        point_count = "2-3"
    elif word_count < 1000:
        point_count = "3-4"
    else:
        point_count = "4-5"
    
    # 添加强制语言指令
    if language.lower() == 'chinese':
        language_instruction = "You must respond completely in Chinese (中文). Do not use any English words."
    elif language.lower() == 'spanish':
        language_instruction = "You must respond completely in Spanish (Español). Do not use any English words."
    else:
        language_instruction = "You must respond completely in English."
    
    return f"""
{role_context}, extract {point_count} key strategic insights and takeaways from the following content.

{language_instruction}

Content:
{text}

Requirements:
- Write entirely in {language}
- Must use numbered format (1. 2. 3.)
- Focus on "what this means" and "why it matters" for your role
- Provide actionable insights and implications
- Each point should be substantial and valuable
- No introductory text, just the numbered insights
- CRITICAL: Your entire response must be in {language} language only
"""

def create_question_prompt(role, goal, language, text):
    """创建问题提示词"""
    
    role_contexts = {
        'student': 'As a student',
        'teacher': 'As a teacher',
        'sales': 'As a sales professional', 
        'customer': 'As a customer',
        'product manager': 'As a product manager',
        'analyst': 'As an analyst',
        'supervisor': 'As a supervisor',
        'reader': 'As a reader',
        'consultant': 'As a consultant',
        'developer': 'As a developer',
        'marketer': 'As a marketer'
    }
    
    role_context = role_contexts.get(role, 'As a reader')
    
    word_count = len(text.split())
    if word_count < 300:
        point_count = "2-3"
    elif word_count < 1000:
        point_count = "3-4"
    else:
        point_count = "4-5"
    
    # 添加强制语言指令
    if language.lower() == 'chinese':
        language_instruction = "You must respond completely in Chinese (中文). Do not use any English words."
    elif language.lower() == 'spanish':
        language_instruction = "You must respond completely in Spanish (Español). Do not use any English words."
    else:
        language_instruction = "You must respond completely in English."
    
    return f"""
{role_context}, generate {point_count} thought-provoking questions about the following content.

{language_instruction}

Content:
{text}

Requirements:
- Write entirely in {language}
- Must use numbered format (1. 2. 3.)
- Ask questions that encourage deeper thinking
- Challenge assumptions and explore implications
- Questions should be relevant for your role perspective
- No introductory text, just the numbered questions
- CRITICAL: Your entire response must be in {language} language only
"""

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)